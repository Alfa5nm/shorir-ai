from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"D:\Aust\SHORIR AI")
OUT = ROOT / "deliverables" / "SHORIR_AI_Technical_Report.docx"
BLUE = RGBColor(20, 103, 216)
INK = RGBColor(23, 32, 51)
MUTED = RGBColor(104, 117, 138)
LIGHT = "E9F2FF"
LINE = "D9E2EE"
FRAMES = ROOT / "deliverables" / "demo-frames"


def set_font(run, size=11, bold=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_heading(doc, value, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(value)
    return p


def add_body(doc, value, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and value.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(value[len(bold_lead):]))
    else:
        set_font(p.add_run(value))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        set_font(p.add_run(item))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.2
        set_font(p.add_run(item))


def add_callout(doc, label, body):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    p_pr.append(shd)
    set_font(p.add_run(label.upper()), 9, True, BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.12)
    p2.paragraph_format.right_indent = Inches(0.12)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    p2_pr = p2._p.get_or_add_pPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:fill"), LIGHT)
    p2_pr.append(shd2)
    set_font(p2.add_run(body), 11, True, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "E8EEF5")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        set_font(p.add_run(header), 9.5, True, INK)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row, widths)):
            cells[idx].width = Inches(width)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(value), 9.5, False, INK)
    return table


def add_figure(doc, path, caption, width=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    set_font(cap.add_run(caption), 9, False, MUTED)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25
for name, size, before, after in (
    ("Heading 1", 17, 18, 8),
    ("Heading 2", 14, 14, 6),
    ("Heading 3", 12, 10, 4),
):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = BLUE if name != "Heading 3" else INK
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("SHORIR AI  |  Technical Report"), 8.5, True, MUTED)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("Team El Bracino | Mindsparks 26 CodeFront Challenge | June 2026"), 8.5, False, MUTED)

# Editorial cover
doc.add_paragraph().paragraph_format.space_after = Pt(72)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("TECHNICAL REPORT"), 10, True, BLUE)
p.paragraph_format.space_after = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("SHORIR AI"), 32, True, INK, "Calibri")
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Bangla-First Privacy-Aware Fitness Intelligence"), 18, False, MUTED, "Calibri")
p.paragraph_format.space_after = Pt(26)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(str(ROOT / "apps/web/public/images/logo_nobg.png"), width=Inches(2.0))
p.paragraph_format.space_after = Pt(38)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Prepared by Team El Bracino"), 12, True, INK)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p2.add_run("Mindsparks 26 CodeFront Challenge | AUST Innovation and Design Club"), 10.5, False, MUTED)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p3.add_run("26 June 2026"), 10.5, False, MUTED)
doc.add_page_break()

add_heading(doc, "Executive Summary", 1)
add_body(doc, "SHORIR AI is a Bangla-first fitness companion that combines browser-side pose analysis, a bilingual exercise library, personalized workout planning, Bangladeshi diet guidance, image-based calorie review, progress diagnostics, and a guided demo flow in one coordinated product.")
add_callout(doc, "Core design decision", "When pose tracking is uncertain, the system pauses and explains why. It prioritizes missed repetitions over false repetitions.")
add_body(doc, "The system uses a React and Vite frontend, an Express API, shared TypeScript contracts, and ports-and-adapters boundaries for MediaPipe, persistence, and AI services. The refreshed interface uses compact navigation, shadcn-style local primitives, tabbed secondary detail, and a white-blue visual system with dark mode. Raw live pose video remains in the browser. Phone video uses encrypted WebRTC, with the API handling temporary signaling metadata.")

add_heading(doc, "1. Problem and Local Context", 1)
add_body(doc, "Generic fitness applications often assume English fluency, prior knowledge of exercise form, stable camera placement, and access to globally common foods. These assumptions reduce usability for beginners in Bangladesh.")
add_bullets(doc, [
    "Beginners need direct setup, movement, depth, and safety cues.",
    "False repetition counts damage trust and can reinforce poor form.",
    "Workout, food, calorie, and progress tools are commonly fragmented.",
    "Live camera upload creates avoidable privacy and bandwidth costs.",
    "Diet guidance should recognize familiar Bangladeshi meals and portions."
])

add_heading(doc, "2. Product Scope", 1)
add_table(doc, ["Module", "Implemented capability"], [
    ("Daily dashboard", "Profile-aware workout plan using goals, schedule, equipment, safety flags, and recent sessions."),
    ("Pose Coach", "Strict squat, push-up, and forward-lunge analysis with calibration and diagnostics."),
    ("Exercise library", "Bilingual guides, filters, demonstrations, safety cues, and live-coach links."),
    ("Diet Chart", "Bangladeshi meal template scaled to profile-based calorie and macro targets."),
    ("Calorie Check", "Desktop upload and QR/phone capture with cautious structured meal review."),
    ("Progress", "Saved sessions, review history, and detection-quality diagnostics."),
    ("Demo", "Interactive guided tour with real app frames, focal highlights, cursor motion, and presenter controls."),
], [1.45, 5.05])
add_figure(doc, FRAMES / "01-dashboard.png", "Figure 1. Coordinated dashboard with light theme and competition branding.")

add_heading(doc, "3. User Experience and Visual System", 1)
add_body(doc, "The release interface uses a semantic white and blue palette, a fully functional dark mode, compact icon navigation, consistent eight-pixel-or-less radii, thin borders, and responsive layouts. The visual direction adapts the reference design's disciplined grid, whitespace, and restrained controls to an operational fitness product.")
add_bullets(doc, [
    "Light and dark themes persist locally and respect the initial system preference.",
    "Mindsparks 26, CodeFront, and AUST IDC assets appear in the shared product frame.",
    "Dense workflows use stable grids and independent scrolling where appropriate.",
    "Mobile navigation prioritizes icons while preserving accessible names and tooltips.",
    "No route is reduced to a marketing page; the first screen remains a usable product."
])

add_heading(doc, "3.1 Demo Sequence", 2)
add_body(doc, "The live product now includes a dedicated /demo route. It replaces static walkthrough panels with an auto-playing guided tour that uses real app frames, focal highlights, cursor movement, captions, and pause/back/next controls.")
add_table(doc, ["Step", "Route", "Purpose"], [
    ("1", "/onboarding", "Create a clean profile."),
    ("2", "/", "Review today's plan and shortcuts."),
    ("3", "/coach?exercise=squat", "Show strict live rep coaching."),
    ("4", "/exercise-library", "Open movement guidance."),
    ("5", "/diet-chart", "Generate local diet guidance."),
    ("6", "/calorie-check", "Review calories from desktop or phone capture."),
    ("7", "/progress", "Close with saved session diagnostics."),
    ("8", "/demo?scene=submission", "Close on the upload-ready submission package."),
], [0.55, 2.15, 3.8])
add_figure(doc, FRAMES / "02-demo.png", "Figure 2. Interactive demo tour with focal highlight, cursor pointer, and presenter controls.")

add_heading(doc, "4. System Architecture", 1)
add_body(doc, "SHORIR AI is a PNPM TypeScript monorepo. Shared contracts keep browser and API wire shapes aligned, while ports isolate external technology choices.")
add_table(doc, ["Layer", "Technology", "Responsibility"], [
    ("Web", "React 19, Vite", "Routes, workflows, local session state, camera UI, and rendering."),
    ("Pose", "MediaPipe Tasks Vision", "Browser-side landmark estimation and local analysis."),
    ("API", "Express", "Profiles, sessions, coach reviews, meal reviews, and signaling."),
    ("Contracts", "TypeScript package", "Shared request, response, and domain shapes."),
    ("Persistence", "Memory / Supabase", "Demo storage or persistent production storage."),
    ("AI", "Stub / Gemini adapter", "Deterministic development output or configured AI review."),
], [1.05, 1.55, 3.9])
add_numbered(doc, [
    "The browser creates a persistent anonymous profile identifier.",
    "Onboarding stores goals, fitness level, equipment, schedule, body metrics, and safety notes.",
    "Pose landmarks are estimated and analyzed locally; only derived session summaries are saved.",
    "Meal images are submitted only for an explicit review request.",
    "Progress aggregates saved sessions and diagnostic metadata for the same profile."
])

add_heading(doc, "5. Pose Detection and Rep Validation", 1)
add_body(doc, "The original angle-threshold approach was replaced by a conservative temporal state machine. A repetition can advance only while the quality gate remains valid.")
add_heading(doc, "5.1 Quality Gate", 2)
add_bullets(doc, [
    "Required landmark confidence average of approximately 0.68 or higher.",
    "Several consecutive valid frames before calibration or repetition logic advances.",
    "Body-scale stability to reject meaningful camera-distance changes.",
    "Landmark jitter checks to reject unstable tracking.",
    "Exercise-specific orientation, side consistency, and activity-region checks.",
    "Immediate repetition reset when quality fails mid-repetition."
])
add_heading(doc, "5.2 Temporal State Machine", 2)
add_body(doc, "A valid repetition follows a stable top posture, controlled descent, confirmed bottom or depth, controlled ascent, and stable top lockout. Minimum duration, bottom confirmation, and cooldown rules prevent rapid or duplicated counts.")
add_figure(doc, FRAMES / "03-coach.png", "Figure 3. Live squat coach with activity stage, compact readout rail, and hidden detail tabs.")

add_heading(doc, "6. Exercise-Specific Rules", 1)
add_table(doc, ["Exercise", "Primary validity rules"], [
    ("Squat", "Shoulder, hip, knee, and ankle visibility; side view; knee-angle depth; standing reset."),
    ("Push-up", "Shoulder, elbow, wrist, hip, and ankle visibility; plank line; wrist alignment; elbow depth; top lockout."),
    ("Forward lunge", "Split stance, alternating balance, controlled knee depth, stable return, and side-aware calibration."),
], [1.35, 5.15])
add_callout(doc, "Depth limitation", "This is monocular tracking. Depth means relative MediaPipe z and body-scale estimation, not physical depth-sensor measurement.")

add_heading(doc, "7. Nutrition and Meal Review", 1)
add_body(doc, "The diet chart calculates an estimated target from age, gender, height, current weight, and target weight when those values are available. It then scales familiar meal templates across breakfast, lunch, and dinner. Without complete body metrics, the system provides a more general fitness-level estimate.")
add_bullets(doc, [
    "Local meal images and portions improve recognizability.",
    "Calorie and macro values are planning estimates, not clinical prescriptions.",
    "The calorie-check workflow supports desktop upload and QR-based phone capture.",
    "The AI response is structured and cautious, with uncertainty and next-step guidance."
])
add_figure(doc, FRAMES / "06-calorie.png", "Figure 4. Calorie-check workflow with desktop and phone-capture entry points.")

add_heading(doc, "8. Privacy, Safety, and Ethics", 1)
add_bullets(doc, [
    "Raw pose frames are not uploaded to the API.",
    "Phone pose video travels through encrypted peer-to-peer WebRTC when available.",
    "The API stores temporary signaling metadata, not continuous video.",
    "Safety flags influence workout guidance and are visible in the interface.",
    "The product does not diagnose injury or replace medical or nutritional professionals.",
    "Users are instructed to stop when movement causes pain."
])

add_heading(doc, "9. Verification", 1)
add_body(doc, "The repository includes unit tests for daily planning, exercise content, camera geometry, coaching language, pose analysis, and progress diagnostics. The release gate runs type checking, linting, unit tests, and production builds across the monorepo.")
add_table(doc, ["Acceptance scenario", "Expected result"], [
    ("Five valid squats", "Exactly five repetitions; no duplicates."),
    ("Five valid push-ups", "Exactly five repetitions; no extras."),
    ("Arm waving or walking", "Zero repetitions."),
    ("Partial crouches", "Zero squat repetitions."),
    ("Plank shifting or hip sag", "Zero push-up repetitions."),
    ("Camera movement", "Repetition resets and UI explains the pause."),
], [3.0, 3.5])

add_heading(doc, "10. Deployment and Operations", 1)
add_body(doc, "The production configuration builds the full monorepo and starts the Express API. In production, Express serves the compiled Vite application and provides a single-page fallback for deep links. The API health endpoint at /api/health is used for Railway health checks.")
add_bullets(doc, [
    "Runtime: Node.js 20 or later and PNPM 9.",
    "Railway configuration: Railpack builder, root build, root start, health check, and restart policy.",
    "Default demo adapters: in-memory persistence and deterministic stub AI.",
    "Production persistence option: Supabase with environment configuration.",
    "Production AI option: Gemini using API key or application-default credentials."
])

add_heading(doc, "11. Limitations and Future Work", 1)
add_bullets(doc, [
    "Broader real-device calibration is required across body shapes, clothing, lighting, and camera placements.",
    "Monocular body scale cannot provide true physical depth.",
    "WebRTC reliability at scale requires TURN and shared signaling storage.",
    "The memory database resets when the service restarts.",
    "Meal recognition quality depends on the configured AI adapter and image quality.",
    "Future safety guidance should be reviewed with qualified fitness and health professionals."
])

add_heading(doc, "12. Repository and Submission", 1)
add_table(doc, ["Deliverable", "Location"], [
    ("Live website", "https://shorir-ai-production.up.railway.app"),
    ("Demo sequence", "deliverables/SHORIR_AI_Demo_Sequence.md"),
    ("Demo video", "deliverables/SHORIR_AI_Demo_Walkthrough.mp4"),
    ("Project presentation", "deliverables/SHORIR_AI_Project_Presentation.pptx"),
    ("Technical report", "deliverables/SHORIR_AI_Technical_Report.docx"),
    ("Source code archive", "deliverables/SHORIR_AI_Source_Code.zip"),
], [1.8, 4.7])
add_body(doc, "Source code is organized under apps/web, apps/api, packages/contracts, packages/content, supabase, docs, and delegated contributor packages. Secrets, dependency folders, generated builds, and local environment files are excluded from the source archive.")

OUT.parent.mkdir(parents=True, exist_ok=True)
for index, inline_shape in enumerate(doc.inline_shapes, start=1):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", [
        "SHORIR AI product logo",
        "SHORIR AI coordinated dashboard",
        "SHORIR AI demo sequence",
        "SHORIR AI live pose coach",
        "SHORIR AI calorie check"
    ][index - 1])
doc.save(OUT)
print(OUT)
